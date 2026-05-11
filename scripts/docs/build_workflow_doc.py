from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


FONT = "MW Sans"
TEXT = RGBColor(0x11, 0x11, 0x11)
LINK = "0563C1"
INLINE_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)|\*\*([^*]+)\*\*|`([^`]+)`")


def set_run_font(run, *, size=None, bold=None, italic=None, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), LINK)
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    r_fonts.set(qn("w:eastAsia"), FONT)
    r_pr.append(r_fonts)

    new_run.append(r_pr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_markdown_inline(paragraph, text, size=10.5):
    idx = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > idx:
            run = paragraph.add_run(text[idx:match.start()])
            set_run_font(run, size=size, color=TEXT)
        if match.group(1) and match.group(2):
            add_hyperlink(paragraph, match.group(1), match.group(2))
        elif match.group(3):
            run = paragraph.add_run(match.group(3))
            set_run_font(run, size=size, bold=True, color=TEXT)
        elif match.group(4):
            run = paragraph.add_run(match.group(4))
            set_run_font(run, size=size, italic=True, color=TEXT)
        idx = match.end()
    if idx < len(text):
        run = paragraph.add_run(text[idx:])
        set_run_font(run, size=size, color=TEXT)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.18

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10.5)

    if "AHA Title" not in doc.styles:
        style = doc.styles.add_style("AHA Title", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(24)
        style.font.bold = True
        style.font.color.rgb = TEXT

    if "AHA Heading" not in doc.styles:
        style = doc.styles.add_style("AHA Heading", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(16)
        style.font.bold = True
        style.font.color.rgb = TEXT

    if "AHA Subheading" not in doc.styles:
        style = doc.styles.add_style("AHA Subheading", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(12)
        style.font.bold = True
        style.font.color.rgb = TEXT

    core = doc.core_properties
    core.title = "AHA Design to Code Workflow"
    core.subject = "Working draft"


def add_page_break_section(doc):
    new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    new_section.top_margin = Inches(0.9)
    new_section.bottom_margin = Inches(0.9)
    new_section.left_margin = Inches(1.0)
    new_section.right_margin = Inches(1.0)


def add_title(doc, title, subtitle):
    p = doc.add_paragraph(style="AHA Title")
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(title)
    set_run_font(run, size=24, bold=True, color=TEXT)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(14)
    run = p2.add_run(subtitle)
    set_run_font(run, size=11, italic=True, color=TEXT)


def add_heading(doc, text):
    p = doc.add_paragraph(style="AHA Heading")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    set_run_font(run, size=16, bold=True, color=TEXT)


def add_subheading(doc, text):
    p = doc.add_paragraph(style="AHA Subheading")
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=12, bold=True, color=TEXT)


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.18
    add_markdown_inline(p, text)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    add_markdown_inline(p, text)


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    add_markdown_inline(p, text)


def build_doc(source_path: Path, output_path: Path):
    lines = source_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)

    first_heading_written = False
    title_written = False

    for line in lines:
        if not line.strip():
            continue
        if line.startswith("# "):
            add_title(doc, line[2:].strip(), "Working draft")
            title_written = True
            continue
        if line.startswith("## "):
            if first_heading_written:
                add_page_break_section(doc)
            add_heading(doc, line[3:].strip())
            first_heading_written = True
            continue
        if line.startswith("### "):
            add_subheading(doc, line[4:].strip())
            continue
        if re.match(r"^\d+\.\s+", line):
            add_number(doc, re.sub(r"^\d+\.\s+", "", line).strip())
            continue
        if line.startswith("- "):
            add_bullet(doc, line[2:].strip())
            continue
        add_body(doc, line.strip())

    if not title_written:
        raise RuntimeError("Source markdown missing title")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


if __name__ == "__main__":
    repo = Path(__file__).resolve().parents[2]
    source = repo / "reference" / "docs" / "aha-design-to-code-workflow-v1.md"
    output = repo / ".artifacts" / "build" / "docs" / "AHA Design to Code Workflow.docx"
    build_doc(source, output)
    print(output)
